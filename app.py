import streamlit as st
import os
import tempfile
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt

def convert_pdf_to_docx(pdf_file, output_path):
    """Converts PDF to DOCX retaining layout."""
    cv = Converter(pdf_file)
    cv.convert(output_path, start=0, end=None)
    cv.close()

def change_font(doc_path, font_name, font_size):
    """Opens a DOCX file and forces specific font family and size."""
    doc = Document(doc_path)
    
    # 1. Update the 'Normal' style (helps with future editing in the doc)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    
    # 2. Force apply to all existing paragraph runs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            
    # 3. Force apply to all tables (common in PDF conversions)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        
    doc.save(doc_path)

# --- UI Setup ---
st.set_page_config(page_title="PDF to Cursive Converter", layout="centered")
st.title("📝 PDF to Cursive Word Doc")
st.write("Upload a PDF to convert it to an editable Word document with specific handwriting fonts.")

# 1. File Upload
uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

# 2. Configuration Columns
col1, col2 = st.columns(2)

with col1:
    # Set "Twinkl Cursive Looped" as the first item (default)
    font_option = st.selectbox(
        "Select Output Font",
        (
            "Twinkl Cursive Looped", 
            "Twinkl Cursive Unlooped", 
            "Arial", 
            "Calibri", 
            "Times New Roman",
            "Comic Sans MS"
        )
    )

with col2:
    # Font Size (Default set to 11 for better readability of cursive)
    font_size = st.number_input(
        "Font Size (pt)", 
        min_value=8, 
        max_value=72, 
        value=12, 
        step=1
    )

if uploaded_file is not None:
    # Create a temporary directory to store files during processing
    with tempfile.TemporaryDirectory() as temp_dir:
        
        pdf_path = os.path.join(temp_dir, "input.pdf")
        docx_path = os.path.join(temp_dir, "converted_output.docx")
        
        # Save uploaded PDF to temp location
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        if st.button("Convert to Word"):
            with st.spinner("Processing... extracting text and applying styles."):
                try:
                    # Step A: Convert PDF to DOCX
                    convert_pdf_to_docx(pdf_path, docx_path)
                    
                    # Step B: Apply Font Configuration
                    change_font(docx_path, font_option, font_size)
                    
                    # Step C: Read file back for download
                    with open(docx_path, "rb") as f:
                        file_data = f.read()
                    
                    st.success(f"Converted successfully to {font_option} at {font_size}pt!")
                    
                    # Step D: Download Button
                    st.download_button(
                        label="Download Word Doc",
                        data=file_data,
                        file_name="cursive_worksheet.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.info(f"**Note:** Ensure your computer has the font **'{font_option}'** installed to view the file correctly.")
                    
                except Exception as e:

                    st.error(f"An error occurred: {e}")
